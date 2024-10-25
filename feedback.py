import os
import docx
import math
import pandas as pd


CURRENT_DIRECTORY = os.path.dirname(os.path.realpath(__file__)).replace("\\", "/")
OUTPUT_DIRECTORY = f"{CURRENT_DIRECTORY}/feedback"
FEEDBACK_FILE_NAME = "feedback.csv"
FEEDBACK_FILE_PATH = f"{CURRENT_DIRECTORY}/{FEEDBACK_FILE_NAME}"


def get_column_index(column):
    # 26 number system where [A...Z] is mapped to [1...26]
    decimal_value = 0
    for idx in reversed(range(0, len(column))):
        decimal_value += (ord(column[idx]) - ord("A") + 1) * math.pow(26, len(column) - idx - 1)

    return int(decimal_value - 1)  # the index is the decimal value minus 1


NUM_TABLE_COLUMNS = 3
EMAIL_IDX = get_column_index("B")
NAME_IDX = get_column_index("C")
START_COLUMN_IDX = NAME_IDX + 1
EMAIL_TITLE = "Имейл/телефон"
NAME_TITLE = "Две имена"
FEEDBACK_TITLE = "Обратна връзка"


def try_create_doc(csv_data, team_name, email_column_data: list, name_column_data: list, feedback_column_data: list, file_path: str):
    doc = docx.Document()
    doc.add_heading(team_name, 0)

    num_table_rows = len(email_column_data) + 1
    table = doc.add_table(rows=num_table_rows, cols=NUM_TABLE_COLUMNS)
    table.cell(0, 0).text = EMAIL_TITLE
    table.cell(0, 1).text = NAME_TITLE
    table.cell(0, 2).text = FEEDBACK_TITLE

    for r in range(num_table_rows - 1):
        for c in range(NUM_TABLE_COLUMNS):
            cell = table.cell(r + 1, c)
            if c == 0:
                cell.text = cell_data_to_string(email_column_data[r])
            elif c == 1:
                cell.text = cell_data_to_string(name_column_data[r])
            else:
                cell.text = cell_data_to_string(feedback_column_data[r])

    doc.save(file_path)
    return True


def create_docs():
    if not os.path.exists(OUTPUT_DIRECTORY):
        os.mkdir(OUTPUT_DIRECTORY)

    csv_data = pd.read_csv(FEEDBACK_FILE_PATH)
    email_column_data = get_column_data(csv_data, EMAIL_IDX)
    name_column_data = get_column_data(csv_data, NAME_IDX)
    num_columns = csv_data.shape[1]
    for column_idx in range(START_COLUMN_IDX, num_columns):
        feedback_column_data = get_column_data(csv_data, column_idx)
        team_name = get_column_name(csv_data, column_idx)
        file_name = team_name.replace(" ", "").replace(":", "_")
        file_path = f"{OUTPUT_DIRECTORY}/{file_name}.docx".replace("\\", "/")
        try_create_doc(csv_data, team_name, email_column_data, name_column_data, feedback_column_data, file_path)


def get_column_name(csv_data, column_index: int):
    column_name = csv_data.columns[column_index]
    return column_name


def get_column_data(csv_data, column_index: int):
    column_name = get_column_name(csv_data, column_index)
    column_data = csv_data[column_name].tolist()
    return column_data


def cell_data_to_string(cell_data):
    string_value = str(cell_data)
    if string_value == "nan":
        return ""
    else:
        return string_value


if __name__ == "__main__":
    create_docs()
