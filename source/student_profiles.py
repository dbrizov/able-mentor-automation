import os
import csv
import docx
import math

# If your are using Visual Studio Code there may be a problem with importing the docx lib so you can run " pip install python-docx "

# For the script to work you will need to download the current season's sheet table as a .csv file in the 'file' menu and put it into the same folder as this script.
# Change the name of the .csv file to be student_register.csv
# Then change the targeted table column names such as "O","AF" or "AL" in the functions down below to the coresponding column names in this season's register
# After final a check that everything is right, you can run the script in any terminal or cmd and go to sleep

CURRENT_DIRECTORY = os.path.dirname(
    os.path.realpath(__file__)).replace("\\", "/")
OUTPUT_DIRECTORY = f"{CURRENT_DIRECTORY}/student_profiles"
REGISTER_FILE_NAME = "student_register.csv"
REGISTER_FILE_PATH = f"{CURRENT_DIRECTORY}/{REGISTER_FILE_NAME}"


def get_column_index(column):
    # 26 number system where [A...Z] is mapped to [1...26]
    decimal_value = 0
    for idx in reversed(range(0, len(column))):
        decimal_value += (ord(column[idx]) - ord("A") + 1) * \
            math.pow(26, len(column) - idx - 1)

    return int(decimal_value - 1)  # the index is the decimal value minus 1


CONFIRMED = get_column_index("O")
STUDENT_NAME = get_column_index("AF")
AGE = get_column_index("AL")
CITY = get_column_index("AN")
SCHOOL_NAME = get_column_index("AM")
GRADE = get_column_index("AO")
SCHOOL_INTERESTS = get_column_index("AP")
NON_SCHOOL_INTERESTS = get_column_index("AQ")
SIMILAR_PROGRAMS = get_column_index("AR")
NEED = get_column_index("AS")
ENGLISH_LEVEL = get_column_index("AT")
SPORT = get_column_index("AU")
WHAT_TO_DO_AFTER_SCHOOL = get_column_index("AV")
INTERESTS = get_column_index("AW")
MENTOR_EXPERIENCE = get_column_index("AX")
SKILLS_TO_IMPROVE = get_column_index("AY")
FREE_TIME_ACTIVITIES = get_column_index("AZ")
DIFFICULT_SITUATION = get_column_index("BA")
WHY_APPLY = get_column_index("BB")
IDEA_IN_ABLE_MENTOR = get_column_index("BC")
WANT_TO_CHANGE = get_column_index("BD")
HOURS_PER_WEEK = get_column_index("BE")
PROJECT_WITH_MENTOR = get_column_index("BF")
HEARD_OF_ABLE_MENTOR = get_column_index("BG")

column_titles = {
    CONFIRMED: "Потвърдил участие",
    STUDENT_NAME: "Ученик",
    AGE: "Възраст",
    CITY: "Населено място",
    SCHOOL_NAME: "Училище",
    GRADE: "Завършен клас",
    SCHOOL_INTERESTS: "Интереси, свързани с училище",
    NON_SCHOOL_INTERESTS: "Интереси извън училище",
    SIMILAR_PROGRAMS: "Участвал ли си в други сходни програми, завършил ли си ги и какво си взе от тях?",
    NEED: "ABLE Mentor е напълно безплатна програма с ограничен капацитет. Защо имаш нужда да участваш в нея?",
    ENGLISH_LEVEL: "Ниво на английски език",
    SPORT: "Спорт",
    WHAT_TO_DO_AFTER_SCHOOL: "Какво ще правя след гимназията:",
    INTERESTS: "В кои сфери имаш интерес да се развиваш и в кои по-слаб?",
    MENTOR_EXPERIENCE: "Ментор в каква професионална сфера би бил/а най-полезен/а за теб?",
    SKILLS_TO_IMPROVE: "Кои свои качества искаш да промениш/подобриш?",
    FREE_TIME_ACTIVITIES: "Как се забавляваш в свободното си време?",
    DIFFICULT_SITUATION: "Разкажи ни за трудна ситуация и как си се справил/а?",
    WHY_APPLY: "Защо кандидатстваш в програмата?",
    IDEA_IN_ABLE_MENTOR: "Каква идея искаш да осъществиш в рамките на ABLE Mentor?",
    WANT_TO_CHANGE: "Желая да променя...",
    HOURS_PER_WEEK: "По колко часа седмично би отделял/а на проекта?",
    PROJECT_WITH_MENTOR: "По какъв проект би работил/а със своя ментор?",
    HEARD_OF_ABLE_MENTOR: "Научил/а за ABLE Mentor от?"
}


def try_create_doc(row_data, file_path):
    if row_data[CONFIRMED] != "Да":
        return False

    doc = docx.Document()
    doc.add_heading("ПРОФИЛ НА ТВОЯ УЧЕНИК", 0)

    table = doc.add_table(rows=0, cols=2)

    for idx in range(0, len(row_data)):
        if (idx == CONFIRMED or
            idx == STUDENT_NAME or
                idx not in column_titles):
            continue

        if idx > HEARD_OF_ABLE_MENTOR:
            break

        table_row = table.add_row().cells
        table_row[0].text = column_titles[idx]
        table_row[1].text = row_data[idx]

    doc.save(file_path)
    return True


def create_docs():
    if not os.path.exists(OUTPUT_DIRECTORY):
        os.mkdir(OUTPUT_DIRECTORY)

    doc_counter = 1  # Start a simple counter for numbering the files

    with open(REGISTER_FILE_PATH, encoding="utf-8", mode="r") as fstream:
        reader = csv.reader(fstream, delimiter=',', quotechar='"')

        for idx, row in enumerate(reader):
            if idx == 0:
                continue

            student_name = row[STUDENT_NAME]
            file_path = f"{OUTPUT_DIRECTORY}/{doc_counter}_{student_name}.docx"

            if try_create_doc(row, file_path):
                doc_counter += 1  # Increment the counter only if the document is created successfully


if __name__ == "__main__":
    create_docs()
