import io
import os
import csv
import docx
import matplotlib.pyplot as plt
import numpy


CURRENT_DIRECTORY = os.path.dirname(os.path.realpath(__file__)).replace("\\", "/")
OUTPUT_DIRECTORY_PRESENT = f"{CURRENT_DIRECTORY}/student_self_evaluations_present"
OUTPUT_DIRECTORY_ONLINE = f"{CURRENT_DIRECTORY}/student_self_evaluations_online"
RESPONSES_FILE_NAME_PRESENT = "self_evaluations_responses_beginning_present.csv"
RESPONSES_FILE_NAME_ONLINE = "self_evaluations_responses_beginning_online.csv"
RESPONSES_FILE_PATH_PRESENT = f"{CURRENT_DIRECTORY}/{RESPONSES_FILE_NAME_PRESENT}"
RESPONSES_FILE_PATH_ONLINE = f"{CURRENT_DIRECTORY}/{RESPONSES_FILE_NAME_ONLINE}"

INDEX_IMPORTANT_THINGS = 34
INDEX_STUDENT_NAME = 35
INDEX_SEND_TO_MENTOR = 36

column_titles = [None] * 37
column_titles[0] = "Timestamp"
column_titles[1] = "Email Address"

column_titles[2] = "Убедителност при говорене със събеседник"  # [Важност]
column_titles[3] = "Убедителност при говорене със събеседник"  # [Увереност]
column_titles[4] = "Умение за говорене пред публика"  # [Важност]
column_titles[5] = "Умение за говорене пред публика"  # [Увереност]
column_titles[6] = "Правилна онлайн комуникация по имейл"  # [Важност]
column_titles[7] = "Правилна онлайн комуникация по имейл"  # [Увереност]
column_titles[8] = "Получаване, даване на конструктивна обратна връзка"  # [Важност]
column_titles[9] = "Получаване, даване на конструктивна обратна връзка"  # [Увереност]
column_titles[10] = "Активно слушане"  # [Важност]
column_titles[11] = "Активно слушане"  # [Увереност]

column_titles[12] = "Успешна и ефективна работа в екип"  # [Важност]
column_titles[13] = "Успешна и ефективна работа в екип"  # [Увереност]
column_titles[14] = "Генериране на идеи"  # [Важност]
column_titles[15] = "Генериране на идеи"  # [Увереност]
column_titles[16] = "Структуриране на проект"  # [Важност]
column_titles[17] = "Структуриране на проект"  # [Увереност]
column_titles[18] = "Търсене и анализ на информация"  # [Важност]
column_titles[19] = "Търсене и анализ на информация"  # [Увереност]
column_titles[20] = "Изготвяне на презентация"  # [Важност]
column_titles[21] = "Изготвяне на презентация"  # [Увереност]
column_titles[22] = "Привличане на съмишленици, партньори, спонсори за реализиране на проект"  # [Важност]
column_titles[23] = "Привличане на съмишленици, партньори, спонсори за реализиране на проект"  # [Увереност]

column_titles[24] = "Поставяне на SMART цели"  # [Важност]
column_titles[25] = "Поставяне на SMART цели"  # [Увереност]
column_titles[26] = "Приоритизиране на задачи"  # [Важност]
column_titles[27] = "Приоритизиране на задачи"  # [Увереност]
column_titles[28] = "Ефективно планиране и управление на времето"  # [Важност]
column_titles[29] = "Ефективно планиране и управление на времето"  # [Увереност]
column_titles[30] = "Вътрешна мотивация"  # [Важност]
column_titles[31] = "Вътрешна мотивация"  # [Увереност]
column_titles[32] = "Ясна представа за следващите 5 години"  # [Важност]
column_titles[33] = "Ясна представа за следващите 5 години"  # [Увереност]

column_titles[34] = "Напиши най - важните 3 неща за теб, които искаш да постигнеш с участието си в програмата ABLE Mentor:"
column_titles[35] = "Казвам се..."
column_titles[36] = "Съгласен съм резултатът от моя тест да бъде даден за информация на моя ментор"


def create_bar_chart(title, bar_labels, data):
    x = numpy.arange(len(bar_labels))  # the label locations
    width = 0.2  # the width of the bars
    multiplier = 0

    fig, ax = plt.subplots(layout="constrained")
    fig.set_figwidth(7)
    fig.set_figheight(5.3)

    for attribute, measurement in data.items():
        offset = width * multiplier
        rects = ax.bar(x + offset, measurement, width, label=attribute)
        ax.bar_label(rects, padding=3)
        multiplier += 1

    ax.set_title(title)
    ax.set_xticks(x + width, bar_labels)
    ax.legend(loc="upper right", ncols=2)
    ax.set_ylim(0, 11.9)

    fig_bytes = io.BytesIO()
    plt.savefig(fig_bytes)
    plt.close()  # close current figure to free memory
    fig_bytes.seek(0)
    return fig_bytes


def try_create_doc(student_name, row_data, file_path):
    doc = docx.Document()

    # heading
    doc.add_heading(f"Резултат от тест за самооценка\n({student_name})", 0)

    # 3 important things
    p = doc.add_paragraph()
    p.add_run(f"{column_titles[INDEX_IMPORTANT_THINGS]}").bold = True
    doc.add_paragraph(f'"{row_data[INDEX_IMPORTANT_THINGS]}"')

    # communication
    bar_labels = [x.replace(" ", "\n") for x in column_titles[2:12:2]]
    data = {
        "Важност - начало": [int(x) for x in row_data[2:12:2]],
        "Увереност - начало": [int(x) for x in row_data[3:12:2]],
    }

    png = create_bar_chart("Комуникация", bar_labels, data)
    doc.add_picture(png)
    # plt.savefig(file_path.replace(".docx", ".png"))  # debug png

    # business skills
    bar_labels = [x.replace(" ", "\n") for x in column_titles[12:24:2]]
    data = {
        "Важност - начало": [int(x) for x in row_data[12:24:2]],
        "Увереност - начало": [int(x) for x in row_data[13:24:2]],
    }

    png = create_bar_chart("Бизнес умения", bar_labels, data)
    doc.add_picture(png)

    # personal effectiveness
    bar_labels = [x.replace(" ", "\n") for x in column_titles[24:34:2]]
    data = {
        "Важност - начало": [int(x) for x in row_data[24:34:2]],
        "Увереност - начало": [int(x) for x in row_data[25:34:2]],
    }

    png = create_bar_chart("Лична ефективност", bar_labels, data)
    doc.add_picture(png)

    doc.save(file_path)
    return True


def create_docs(output_directory, responses_file_path):
    if not os.path.exists(output_directory):
        os.mkdir(output_directory)

    with open(responses_file_path, encoding="utf-8", mode="r") as fstream:
        reader = csv.reader(fstream, delimiter=',', quotechar='"')

        for idx, row in enumerate(reader):
            if idx == 0:
                continue  # skip first row

            student_name = row[INDEX_STUDENT_NAME].replace("/", "").strip()
            file_path = ""
            if row[INDEX_SEND_TO_MENTOR].startswith("Не"):
                file_path = f"{output_directory}/{student_name}_НЕ.docx".replace("\\", "/")
            else:
                file_path = f"{output_directory}/{student_name}.docx".replace("\\", "/")

            try_create_doc(student_name, row, file_path)


if __name__ == "__main__":
    create_docs(OUTPUT_DIRECTORY_PRESENT, RESPONSES_FILE_PATH_PRESENT)
    create_docs(OUTPUT_DIRECTORY_ONLINE, RESPONSES_FILE_PATH_ONLINE)
